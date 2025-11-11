import io
from datetime import date
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ---------- Helpers ----------
def rupiah(n: int) -> str:
    return f"{n:,}".replace(",", ".")

def parse_int(s: str) -> int:
    if not s:
        return 0
    s = s.replace(".", "").replace(",", "").strip()
    return int(s) if s.isdigit() else 0

def add_row(table, left: str, right: str, right_align=True):
    r = table.add_row().cells
    r[0].text = left
    p = r[1].paragraphs[0]
    run = p.add_run(right)
    if right_align:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return r

def small_para(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0

def set_normal_style(doc: Document):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(11)

def new_table(doc, cols=2, col2_width_cm=5.0):
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.autofit = False
    # kolom nominal agak sempit, biar hemat ruang
    tbl.columns[0].width = Cm(14.5 - col2_width_cm)
    tbl.columns[1].width = Cm(col2_width_cm)
    for cell in tbl.rows[0].cells:
        for p in cell.paragraphs:
            small_para(p)
    return tbl

def add_section_title(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].bold = True
    small_para(p)

# ---------- DOCX builder ----------
def build_docx(
    bulan_tahun: str,
    nama_bulan_sebelumnya: str,
    saldo_awal: int,
    kencleng_kali: int,
    kencleng_total: int,
    income_custom: list,         # list of (desc:str, amount:int)
    rt_breakdown: list,          # list panjang 5 angka int
    khotib_kali: int,
    honor_khotib: int,
    honor_marbot: int,
    listrik_kali: int,
    bayar_listrik: int,
    expense_custom: list,        # list of (desc, amount)
    ttd_kota_tanggal: str,
    nama_ketua: str,
    nama_bendahara: str
):
    doc = Document()
    set_normal_style(doc)

    # Margin tipis biar muat 1 lembar
    sec = doc.sections[0]
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(1.6)
    sec.right_margin = Cm(1.6)

    # Judul
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Laporan Keuangan DKM Sirojul Huda")
    r.bold = True; r.font.size = Pt(12)
    small_para(p)

    p = doc.add_paragraph("Aspol Sukamiskin Bandung")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    small_para(p)

    p = doc.add_paragraph(f"Bulan {bulan_tahun}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    small_para(p)

    # Saldo sebelumnya
    add_section_title(doc, f"Saldo {nama_bulan_sebelumnya}")
    p = doc.add_paragraph(rupiah(saldo_awal))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    small_para(p)

    # Pemasukan
    doc.add_paragraph()  # spasi tipis
    add_section_title(doc, "Pemasukan")
    tbl_in = new_table(doc, col2_width_cm=4.3)

    # Kencleng (fixed)
    add_row(tbl_in, f"Kencleng Jumat ({kencleng_kali}x)", rupiah(kencleng_total))

    total_pemasukan = kencleng_total

    # Custom income (skip yang kosong)
    for desc, amt in income_custom:
        desc = (desc or "").strip()
        amt_i = parse_int(amt) if isinstance(amt, str) else int(amt or 0)
        if desc and amt_i:
            add_row(tbl_in, desc, rupiah(amt_i))
            total_pemasukan += amt_i

    # Total pemasukan
    add_row(tbl_in, "Total Pemasukan", rupiah(total_pemasukan))

    # Rincian RW 07 (opsional; hanya tampil bila ada salah satu diisi)
    if any(x > 0 for x in rt_breakdown):
        doc.add_paragraph()
        add_section_title(doc, "Rincian Infaq Warga RW 07 (berdasarkan RT)")
        tbl_rt = new_table(doc, col2_width_cm=4.0)
        for i, val in enumerate(rt_breakdown, start=1):
            if val > 0:
                add_row(tbl_rt, f"RT {i:02d}", rupiah(val))

    # Pemasukan Kotor
    doc.add_paragraph()
    add_section_title(doc, "Pemasukan Kotor")
    pemasukan_kotor = saldo_awal + total_pemasukan
    p = doc.add_paragraph(rupiah(pemasukan_kotor))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    small_para(p)

    # Pengeluaran
    doc.add_paragraph()
    add_section_title(doc, "Pengeluaran")
    tbl_out = new_table(doc, col2_width_cm=4.3)

    total_pengeluaran = 0

    # Fixed 3 items
    if honor_khotib:
        add_row(tbl_out, f"Honor Khotib ({khotib_kali}x)", rupiah(honor_khotib))
        total_pengeluaran += honor_khotib
    if honor_marbot:
        add_row(tbl_out, "Honor Marbot + Uang Saku", rupiah(honor_marbot))
        total_pengeluaran += honor_marbot
    if bayar_listrik:
        add_row(tbl_out, f"Bayar Listrik ({listrik_kali}x)", rupiah(bayar_listrik))
        total_pengeluaran += bayar_listrik

    # Custom expenses
    for desc, amt in expense_custom:
        desc = (desc or "").strip()
        amt_i = parse_int(amt) if isinstance(amt, str) else int(amt or 0)
        if desc and amt_i:
            add_row(tbl_out, desc, rupiah(amt_i))
            total_pengeluaran += amt_i

    add_row(tbl_out, "Total Pengeluaran", rupiah(total_pengeluaran))

    # Saldo akhir
    doc.add_paragraph()
    add_section_title(doc, f"Saldo {bulan_tahun}")
    saldo_akhir = pemasukan_kotor - total_pengeluaran
    p = doc.add_paragraph(rupiah(saldo_akhir))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    small_para(p)

    # TTD 2 kolom (ada jarak kosong buat tanda tangan)
    doc.add_paragraph()
    tbl_sig = doc.add_table(rows=2, cols=2)
    tbl_sig.autofit = False
    tbl_sig.columns[0].width = Cm(9.5)
    tbl_sig.columns[1].width = Cm(9.5)

    # Baris jabatan + tanggal
    left = tbl_sig.rows[0].cells[0].paragraphs[0]
    left.add_run("Ketua DKM Sirojul Huda")
    small_para(left)

    right = tbl_sig.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(ttd_kota_tanggal)
    small_para(right)

    # Baris nama dengan ruang tanda tangan di atasnya
    # tambahkan beberapa baris kosong tipis untuk ruang paraf
    for _ in range(2):
        doc.add_paragraph("")

    row2 = doc.add_table(rows=1, cols=2).rows[0].cells
    l = row2[0].paragraphs[0]; l.add_run(nama_ketua); small_para(l)
    r = row2[1].paragraphs[0]; r.alignment = WD_ALIGN_PARAGRAPH.RIGHT; r.add_run(nama_bendahara); small_para(r)

    # Baris “Bendahara DKM” di bawah nama kanan (supaya jelas)
    last = doc.add_table(rows=1, cols=2).rows[0].cells
    last[0].paragraphs[0].add_run("")  # kosong kiri
    pr = last[1].paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr.add_run("Bendahara DKM")
    small_para(pr)

    # Export bytes
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, saldo_akhir, pemasukan_kotor, total_pengeluaran

# ---------- UI ----------
st.set_page_config(page_title="Laporan DKM", layout="wide")
st.title("Laporan Keuangan DKM — Generator DOCX (rapi 1 lembar)")

colA, colB, colC = st.columns([1,1,1])

with colA:
    bulan_tahun = st.text_input("Bulan & Tahun (mis. Oktober 2025)", "Oktober 2025")
    nama_bulan_sebelumnya = st.text_input("Nama Bulan Sebelumnya", "September")
    saldo_awal = parse_int(st.text_input(f"Saldo {nama_bulan_sebelumnya}", "4.113.000"))

with colB:
    st.subheader("Pemasukan (fixed & custom)")
    kencleng_kali = st.number_input("Kencleng Jumat — jumlah kali", min_value=0, value=5, step=1)
    kencleng_total = parse_int(st.text_input("Kencleng Jumat — total", "1.316.000"))

    st.markdown("**Custom Pemasukan** (kosongkan bila tidak dipakai)")
    income_custom = []
    for i in range(1, 8):  # 7 baris custom
        c1, c2 = st.columns([2,1])
        with c1:
            dsc = st.text_input(f"Uraian pemasukan {i}", value="" if i>3 else (["Infaq Warga RW 07 (total)","Infaq Ibu Ninin","Infaq Bapak Agus"][i-1] if i<=3 else ""))
        with c2:
            amt = st.text_input(f"Nominal pemasukan {i}", value="" if i>3 else (["1.740.000","100.000","200.000"][i-1] if i<=3 else ""))
        income_custom.append((dsc, amt))

with colC:
    st.subheader("Rincian RW 07 (RT 1–5)")
    rt_breakdown = []
    defaults = ["290.000","340.000","340.000","370.000","400.000"]
    for i in range(5):
        rt_breakdown.append(parse_int(st.text_input(f"RT {i+1:02d}", defaults[i])))

st.markdown("---")
col1, col2 = st.columns(2)
with col1:
    st.subheader("Pengeluaran (fixed)")
    khotib_kali = st.number_input("Khotib — jumlah Jumat", min_value=0, value=5, step=1)
    honor_khotib = parse_int(st.text_input("Honor Khotib — total", "1.000.000"))
    honor_marbot = parse_int(st.text_input("Honor Marbot + Uang Saku", "1.250.000"))
    listrik_kali = st.number_input("Listrik — jumlah kali", min_value=0, value=2, step=1)
    bayar_listrik = parse_int(st.text_input("Bayar Listrik — total", "176.000"))

with col2:
    st.subheader("Pengeluaran (custom)")
    expense_custom = []
    labels = ["Fotocopy","Alat perbaikan kipas","Ongkos bawa drum","Ongkos pembuatan bedug"]
    values = ["5.000","27.500","65.000","400.000"]
    for i in range(1, 9):  # 8 baris custom
        c1, c2 = st.columns([2,1])
        with c1:
            dsc = st.text_input(f"Uraian pengeluaran {i}", value=(labels[i-1] if i<=4 else ""))
        with c2:
            amt = st.text_input(f"Nominal pengeluaran {i}", value=(values[i-1] if i<=4 else ""))
        expense_custom.append((dsc, amt))

st.markdown("---")
colX, colY = st.columns(2)
with colX:
    ttd_kota_tanggal = st.text_input("Kota & tanggal untuk tanda tangan (kanan atas)", "Bandung, 31 Oktober 2025")
with colY:
    nama_ketua = st.text_input("Nama Ketua DKM", "Ali Marga")
    nama_bendahara = st.text_input("Nama Bendahara DKM", "Eneng Nariah")

if st.button("📝 Buat Dokumen Word"):
    bio, saldo_akhir, pk, tp = build_docx(
        bulan_tahun,
        nama_bulan_sebelumnya,
        saldo_awal,
        int(kencleng_kali),
        kencleng_total,
        income_custom,
        rt_breakdown,
        int(khotib_kali),
        honor_khotib,
        honor_marbot,
        int(listrik_kali),
        bayar_listrik,
        expense_custom,
        ttd_kota_tanggal,
        nama_ketua,
        nama_bendahara
    )

    st.success("Berhasil dibuat. Angka otomatis sudah pakai titik & layout dipadatkan agar 1 lembar.")
    st.write(f"**Pemasukan Kotor:** Rp {rupiah(pk)}  •  **Total Pengeluaran:** Rp {rupiah(tp)}  •  **Saldo Akhir:** Rp {rupiah(pk - tp)}")

    st.download_button(
        "⬇️ Download DOCX",
        data=bio.getvalue(),
        file_name=f"Laporan_DKM_{bulan_tahun.replace(' ', '_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
